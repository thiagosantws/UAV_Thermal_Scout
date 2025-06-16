# Centrais Elétricas do Norte do Brasil S.A.
# Universidade Federal do Pará
# Autor: Thiago Carvalho dos Santos

import os
import cv2
import ast
import numpy as np
from thermal import Thermal
import math
from PIL import Image, ImageTk
from PIL.ExifTags import TAGS, GPSTAGS
from datetime import datetime
from io import BytesIO

import tkinter as tk
from tkinter import scrolledtext
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

pasta_imagens = os.path.abspath('images/')
yolov4_custom = os.path.abspath('files/yolov4_customv2.cfg')
yolov4_weights = os.path.abspath('files/yolov4_customv2_last.weights')
nomes_imagens = os.listdir(pasta_imagens)
quant_imagens = len(nomes_imagens)

area_trafo = 10000
area_chave = 2500

#========================================================VARIÁVEIS DE ENTRADA===================================================================#
  
caminho_txt = os.path.abspath('files/parametros.txt')
with open(caminho_txt, 'r') as arquivo:
    linhas = arquivo.readlines()
    gsd_imagem = float(linhas[0].strip().split('=')[1])
    intervalo_chave = ast.literal_eval(linhas[1].strip().split('=')[1])
    intervalo_trafo = ast.literal_eval(linhas[2].strip().split('=')[1])
    filtro_deteccao = float(linhas[3].strip().split('=')[1])
    confianca = float(linhas[4].strip().split('=')[1])

#=======================================================DETECÇÃO DE EQUIPAMENTO=================================================================#

net = cv2.dnn.readNet(yolov4_weights, yolov4_custom) # Passa os parâmetros dos pesos treinados e da arquitetura do YOLOv4

# REDE NEURAL CONVOLUCIONAL
def processamento(imagem):
    caminho_imagem = os.path.abspath(str(f'images/{imagem}')) # Caminho da imagem
    image = cv2.imread(caminho_imagem) # Abre a imagem 
    height, width, _ = image.shape # Extrai a altura e largura da imagem (512x640)
    # Transforma a imagem para o formato da rede neural: normaliza os vetores RGB, redimensiona para o tamanho 416 x 416,
    # Troca o formato RGB para BGR e não permite o corte da imagem no redimensionamento.
    blob = cv2.dnn.blobFromImage(image, 1/255.0, (416, 416), swapRB=True, crop=False) 
    net.setInput(blob) # Carrega a imagem transformada para dentro da rede
    outputs = net.forward(net.getUnconnectedOutLayersNames()) # Extrai os valores da ultima camada da CNN
    boxes = []
    confidences = []
    class_ids = []

    # Laço para extrair os valores de confiança, classe e 
    for output in outputs:
        for detection in output:
            scores = detection[5:]
            class_id = np.argmax(scores)
            confidence = scores[class_id]

            if confidence > confianca:
                center_x = int(detection[0] * width)
                center_y = int(detection[1] * height)
                w = int(detection[2] * width)
                h = int(detection[3] * height)
                x = int(center_x - w / 2)
                y = int(center_y - h / 2)

                boxes.append([x, y, w, h])
                confidences.append(float(confidence))
                class_ids.append(class_id)
    
    indices = cv2.dnn.NMSBoxes(boxes, confidences, 0.5, 0.4)
    
    boxes2 = []
    class_ids2 = []
    confidences2 = []

    for i in indices:
        i = i.item()
        x, y, w, h = boxes[i]
        class_id = class_ids[i]
        confidence = confidences[i]
        
        boxes2.append([x,y,w,h])
        class_ids2.append(class_id)
        confidences2.append(confidence)

    return boxes2, class_ids2, confidences2

#=====================================================LOCALIZAÇÃO DE EQUIPAMENTO================================================================#

# Função para extração das coordenadas de uma imagem
def coordenada(image_path):
     
    def get_geotagging(im):
        exif_data = im._getexif()
        if exif_data is None:
            raise Exception("No EXIF data found in the image.")
        geotag_data = {}
        for tag, value in exif_data.items():
            tag_name = TAGS.get(tag, tag)
            if tag_name == "GPSInfo":
                for t, v in value.items():
                    sub_tag = GPSTAGS.get(t, t)
                    geotag_data[sub_tag] = v
        if not geotag_data:
            raise Exception("No GPS info found in the image.")
        return geotag_data
    
    def dms_to_decimal(dms, direction):
        degrees = dms[0]
        minutes = dms[1]
        seconds = dms[2]
        decimal = degrees + minutes / 60.0 + seconds / 3600.0
        if direction in ['S','W']:
            decimal = -decimal
        return decimal
    
    image_way = os.path.abspath(str(f'images/{image_path}'))
    image = Image.open(image_way)
    geotag_data = get_geotagging(image)

    if geotag_data:
        latitude = geotag_data.get('GPSLatitude')
        longitude = geotag_data.get('GPSLongitude')
        latitude_direction = geotag_data.get('GPSLatitudeRef', 'N')
        longitude_direction = geotag_data.get('GPSLongitudeRef', 'E')

    latitude_decimal = dms_to_decimal(latitude, latitude_direction)
    longitude_decimal = dms_to_decimal(longitude, longitude_direction)

    return latitude_decimal, longitude_decimal

#==============================================================================================================================================#

# Função para cálculo da orientação do drone  
def orientacao(lat1, lon1, lat2, lon2):
    
    lat1_rad = math.radians(lat1) # Conversão de grau para radiano
    lon1_rad = math.radians(lon1) # Conversão de grau para radiano
    lat2_rad = math.radians(lat2) # Conversão de grau para radiano
    lon2_rad = math.radians(lon2) # Conversão de grau para radiano

    delta_lon = lon2_rad - lon1_rad

    # Cálculo da angulação (sentido-horário) do sentido do drone em relação ao Norte
    angle_rad = math.atan2(math.sin(delta_lon), math.cos(lat1_rad) * math.tan(lat2_rad) - math.sin(lat1_rad) * math.cos(delta_lon))

    angle_degrees = math.degrees(angle_rad) # Conversão radiano para grau
    
    # Garante valores de angulação em graus positivos
    if angle_degrees < 0:
        angle_degrees += 360.0
        
    return angle_degrees

def coord_pixel(lat, lon, angle, x, y):

    image_width = 640 # largura da imagem
    image_height = 512 # altura da imagem
    gsd = gsd_imagem # GSD (m/pixel) da imagem
    y = image_height - y - 60 # Transformando a coordenada para o formato do cálculo da orientação

    angle_rad = np.radians(angle)

    delta_x = (x - (image_width / 2)) * gsd
    delta_y = (y - (image_height / 2)) * gsd

    pixel_latitude = lat + (delta_y * np.cos(angle_rad) - delta_x * np.sin(angle_rad)) / 111000
    pixel_longitude = lon + (delta_x * np.cos(angle_rad) + delta_y * np.sin(angle_rad)) / (111000 * np.cos(np.radians(lat)))

    return pixel_latitude, pixel_longitude

def localizar(item, x, y): #item = posição da imagem na lista de imagens da pasta / x e y = coordenadas do pixel de interesse

    if item == 0:
        latitude1, longitude1 = coordenada(nomes_imagens[item])
        latitude2, longitude2 = coordenada(nomes_imagens[1])
    else:
        latitude1, longitude1 = coordenada(nomes_imagens[item-1])
        latitude2, longitude2 = coordenada(nomes_imagens[item])

    angulo = orientacao(latitude1, longitude1, latitude2, longitude2)
    latitudec, longitudec = coordenada(nomes_imagens[item])

    return coord_pixel(latitudec, longitudec, angulo, x, y) # retorna a latitude e longitude do pixel de interesse

#=========================================================LEITURA DE TEMPERATURA==================================================================#

def temperatura(imagem):
    
    so = 'windows'
    if so == 'windows':
        thermal = Thermal(
            dirp_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/windows/release_x64/libdirp.dll'),
            dirp_sub_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/windows/release_x64/libv_dirp.dll'),
            iirp_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/windows/release_x64/libv_iirp.dll'),
            exif_filename = os.path.abspath('thermal_parser/plugins/exiftool-12.35.exe'),
            dtype = np.float32,
            )
    elif so == 'linux':
        thermal = Thermal(
            dirp_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/linux/release_x64/libdirp.so'),
            dirp_sub_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/linux/release_x64/libv_dirp.so'),
            iirp_filename = os.path.abspath('thermal_parser/plugins/dji_thermal_sdk_v1.1_20211029/linux/release_x64/libv_iirp.so'),
            exif_filename = None,
            dtype = np.float32,
            )
    
    caminho_imagem = os.path.abspath(f'images/{imagem}')
    temperature = thermal.parse_dirp2(image_filename=caminho_imagem, m2ea_mode = True)
    assert isinstance(temperature, np.ndarray)

    return temperature

#=================================================================================================================================================#

img_chaves = [] #lista de imagens com chaves detectadas
img_trafos = [] #lista de imagens com trafos detectados
box_chaves = [] #lista que armazena as caixas delimitadoras das chaves seccionadoras
box_trafos = [] #lista que armazena as caixas delimitadoras dos transformadores

def process(a1,a2):
    
    img_chaves.clear()
    img_trafos.clear()
    box_chaves.clear()
    box_trafos.clear()
    limpar()
    caixa_text.config(state=tk.NORMAL)
    combo.set('Selecione')
    for img in range(a1-1,a2):
        info_img = processamento(nomes_imagens[img])
        if info_img == ([],[],[]):
            texto = str(f'Nenhum equipamento detectado em {nomes_imagens[img]}\n')
            caixa_text.insert(tk.INSERT, texto)
        else:
            if info_img[1].count(0) != 0:
                box_aux = []
                img_chaves.append(nomes_imagens[img])
                for indice, elemento in enumerate(info_img[1]):
                    if elemento == 0:
                        box_aux.append(info_img[0][indice])
                box_chaves.append(box_aux)
                texto = str(f'{info_img[1].count(0)} chave(s) detectada(s) em {nomes_imagens[img]}\n')
                caixa_text.insert(tk.INSERT, texto)
            if info_img[1].count(1) != 0:
                box_aux = []
                img_trafos.append(nomes_imagens[img])
                for indice, elemento in enumerate(info_img[1]):
                    if elemento == 1:
                        box_aux.append(info_img[0][indice])
                box_trafos.append(box_aux)
                texto = str(f'{info_img[1].count(1)} trafo(s) detectado(s) em {nomes_imagens[img]}\n')
                caixa_text.insert(tk.INSERT, texto)
    
    caixa_text.config(state=tk.DISABLED)

def process_imagens():
    try:
        texto = texto_intervalo.get()
        if texto == "":
            process(0,quant_imagens)
        else:
            partes = texto.split('-')
            a1 = int(partes[0])
            a2 = int(partes[1])
            process(a1,a2)
        filtragem()
    except Exception as e:
        messagebox.showerror("Erro de Execução", str(e))

#===================================================================================================================================#

def limite_coord(x,y,w,h):
    x = 0 if x < 0 else (640 if x > 640 else x)
    y = 0 if y < 0 else (512 if y > 512 else y)
    if x+w > 640: w = 640 - x
    if y+h > 512: h = 512 - y
    return x, y, w, h

def info_temp(image, classe, indice):

    temp = np.copy(temperatura(image))
    if classe == 0:
        ind = img_chaves.index(image)
        x, y, w, h = box_chaves[ind][indice]
        x, y, w, h = limite_coord(x,y,w,h)
    elif classe == 1:
        ind = img_trafos.index(image)
        x, y, w, h = box_trafos[ind][indice]
        x, y, w, h = limite_coord(x,y,w,h)
    
    arr = np.empty((h,w))
    for lar in range(w):
        for alt in range(h):
            arr[alt,lar] = temp[y+alt,x+lar]
    
    # Leitura da coordenada da máxima e mínima temperatura
    amax = np.argmax(arr)
    amin = np.argmin(arr)
    coordmax = np.unravel_index(amax, arr.shape)
    coordmin = np.unravel_index(amin, arr.shape)

    # Mapeando coordenada de máxima e mínima temperatura na imagem matriz
    ymax = coordmax[0] + y
    xmax = coordmax[1] + x
    ymin = coordmin[0] + y
    xmin = coordmin[1] + x

    # Valores de máxima, média e mínima temperatura
    tmax = np.max(arr)
    tmin = np.min(arr)
    tmed = np.mean(arr)

    return tmax, tmin, tmed, ymax, xmax, ymin, xmin

#===================================================================================================================================#

def centro(image, classe, indice):

    if classe == 0:
        ind = img_chaves.index(image)
        x, y, w, h = box_chaves[ind][indice]
    elif classe == 1:
        ind = img_trafos.index(image)
        x, y, w, h = box_trafos[ind][indice]
    
    x_center = x + (w/2)
    y_center = y + (h/2)
    
    return x_center, y_center

#===================================================================================================================================#

# Função para detecção de equipamento
def equipamento(image, classe, indice):
    
    item = nomes_imagens.index(image)
    xc, yc = centro(image, classe, indice)
    lat, lon = localizar(item, xc, yc)
    equipamentos = [] # Matriz que recebe | Nome Operacional | Fase | Latitude | Longitude
    
    #PARA CHAVES
    if classe == 0:
        caminho = os.path.abspath('files/chaves.txt')
        with open(caminho, 'r') as arquivo:
            for linha in arquivo:
                elemento = linha.strip().split()
                equipamentos.append(elemento)
    #PARA TRAFO
    elif classe == 1:
        caminho = os.path.abspath('files/trafos.txt')
        with open(caminho, 'r') as arquivo:
            for linha in arquivo:
                elemento = linha.strip().split()
                equipamentos.append(elemento)

    matriz = []
    for coord in range(len(equipamentos)):
        x1 = equipamentos[coord][2]
        y1 = equipamentos[coord][3]
        matriz.append([x1,y1])

    # Função para calcular a distância euclidiana entre dois pontos (x1, y1) e (x2, y2)
    def distancia_euclidiana(x1, y1, x2, y2):
        x1, y1, x2, y2 = float(x1), float(y1), float(x2), float(y2)
        return math.sqrt((x1 - x2) ** 2 + (y1 - y2) ** 2)

    # Função para encontrar a linha mais próxima na matriz
    def linha_mais_proxima(matriz, x, y):
        melhor_linha = None
        menor_distancia = float('inf')  # Inicialize com um valor muito grande
        for i, linha in enumerate(matriz):
            xn, yn = linha
            d = distancia_euclidiana(x, y, xn, yn)
            if d < menor_distancia:
                menor_distancia = d
                melhor_linha = i
        return melhor_linha, menor_distancia

    # Encontre a linha mais próxima na matriz
    linha_equip, distancia = linha_mais_proxima(matriz, lat, lon)
    if distancia > filtro_deteccao:
        nome_equipo = ""
        fase_equipo = ""
    else:
        nome_equipo = equipamentos[linha_equip][0]
        fase_equipo = equipamentos[linha_equip][1]

    return nome_equipo, fase_equipo

#===================================================================================================================================#

def limpar():

    caixa_text.config(state=tk.NORMAL)
    caixa_text.delete(1.0, tk.END)
    combo2["values"] = []
    combo2.set('Selecione')
    combo.set('Selecione')
    text_imagem["text"] = "______________"
    rotulo.config(image=imagem_tk)
    rotulo.image = imagem_tk
    caixa_text.config(state=tk.DISABLED)
    global imagem_atual; imagem_atual = ""
    
def select_combo(arg):
    
    if combo.get() == 'Chave Seccionadora':
        combo2.set('Selecione')
        combo2["values"] = []
        combo2["values"] = img_chaves
    elif combo.get() == 'Transformador':
        combo2.set('Selecione')
        combo2["values"] = []
        combo2["values"] = img_trafos

#=======================================================ANÁLISE DAS IMAGENS==========================================================#

def filtragem(): # Filtragem

    def box_area(box_equipo, classe):
        
        aux_area = False
        if classe == 0:
            _, _, w, h = box_equipo
            if w*h > area_chave:
                aux_area = True
        elif classe == 1:
            _, _, w, h = box_equipo
            if w*h > area_trafo:
                aux_area = True
        return aux_area

    img_chaves_aux = []
    img_trafos_aux = []
    box_chaves_aux = []
    box_trafos_aux = []

    #caixa_text.config(state=tk.NORMAL)
    for i in range(len(img_chaves)): # percorre o indice de todas as imagens com chaves
        box_aux2 = []
        for j in range(len(box_chaves[i])): # for: valor j percorre indice de todas as caixas de uma imagem
            nome, _ = equipamento(img_chaves[i],0,j) # verifica se a caixa está perto de um equipamento
            if nome != "" and box_area(box_chaves[i][j],0) == True: # verifica se a caixa está perto de um equipamento
                box_aux2.append(box_chaves[i][j]) # se estiver perto, adiciona a caixa na lista auxiliar
            #else:
                #caixa_text.insert(tk.INSERT, str(f'Falsa detecção excluída em {img_chaves[i]}\n'))
        if len(box_aux2) != 0: # verifica se há caixas (depois de filtrar) para adicinar o nome da imagem
            box_chaves_aux.append(box_aux2) # adiciona as caixas da imagem na lista final
            img_chaves_aux.append(img_chaves[i])
    for k in range(len(img_trafos)):
        box_aux2 = []
        for l in range(len(box_trafos[k])):
            nome, _ = equipamento(img_trafos[k],1,l)
            if nome != "" and box_area(box_trafos[k][l],1) == True:
                box_aux2.append(box_trafos[k][l])
            #else:
                #caixa_text.insert(tk.INSERT, str(f'Falsa detecção excluída em {img_trafos[k]}\n'))
        if len(box_aux2) != 0:
            box_trafos_aux.append(box_aux2)
            img_trafos_aux.append(img_trafos[k])
    #caixa_text.config(state=tk.DISABLED)

    img_chaves.clear()
    img_trafos.clear()
    box_chaves.clear()
    box_trafos.clear()

    img_chaves[:] = img_chaves_aux
    img_trafos[:] = img_trafos_aux
    box_chaves[:] = box_chaves_aux
    box_trafos[:] = box_trafos_aux

#====================================================================================================================================#

def detector():

    equip_analis = []

    chaves_mono = []
    trafos_mono = []
    caixa_text.config(state=tk.NORMAL)
    if img_chaves == [] and img_trafos == []:
        caixa_text.insert(tk.INSERT, 'Sem equipamento detectado.\n')
    elif combo.get() == 'Selecione':
        caixa_text.insert(tk.INSERT, 'Selecione um equipamento.\n')
    else:
        caixa_text.config(state=tk.NORMAL)
        #PARA CHAVE
        aux01 = False
        if combo.get() == 'Chave Seccionadora' and img_chaves != []:
            aux01 = True
            for i in range(len(img_chaves)):
                for j in range(len(box_chaves[i])):
                    nome, fase = equipamento(img_chaves[i], 0, j)
                    tmax, _, _, _, _, _, _ = info_temp(img_chaves[i], 0, j)
                    box_aux = [nome, fase, tmax]
                    chaves_mono.append(box_aux)

        #PARA TRAFO
        elif combo.get() == 'Transformador' and img_trafos != []:
            aux01 = True
            for i in range(len(img_trafos)):
                for j in range(len(box_trafos[i])):
                    nome, fase = equipamento(img_trafos[i], 1, j)
                    tmax, _, _, _, _, _, _ = info_temp(img_trafos[i], 1, j)
                    box_aux = [nome, fase, tmax]
                    trafos_mono.append(box_aux)
        
        if aux01 == True:
            chaves_monoc = calculo_medias(chaves_mono)
            trafos_monoc = calculo_medias(trafos_mono)
            chaves_dif = trifasico(chaves_monoc)
            trafos_dif = trifasico(trafos_monoc)

            # CLASSIFICAÇÃO DE EQUIPAMENTO
            if combo.get() == 'Chave Seccionadora':
                for name, dif in chaves_dif.items():
                    if dif[0] == 0:
                        texto = str(f'{name} | 1 fase detec | Class: Indefinido')
                    elif dif[0] > 0 and dif[0] <= intervalo_chave[0]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Limite Normal')
                    elif dif[0] > intervalo_chave[0] and dif[0] <= intervalo_chave[1]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Baixo Risco')
                    elif dif[0] > intervalo_chave[1] and dif[0] <= intervalo_chave[2]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Médio Risco')
                    elif dif[0] > intervalo_chave[2] and dif[0] <= intervalo_chave[3]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Alto Risco')
                    elif dif[0] > intervalo_chave[3]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Alto Risco Iminente')
                    caixa_text.insert(tk.INSERT, str(f'{texto}.\n'))
                    partes = texto.split(' | ')
                    equip_analis.append([name,dif[0],dif[1],partes[2].upper(),'chave'])
                    
            elif combo.get() == 'Transformador':
                for name, dif in trafos_dif.items():
                    if dif[0] == 0:
                        texto = str(f'{name} | 1 fase detec | Class: Indefinido')
                    elif dif[0] > 0 and dif[0] <= intervalo_trafo[0]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Limite Normal')
                    elif dif[0] > intervalo_trafo[0] and dif[0] <= intervalo_trafo[1]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Baixo Risco')
                    elif dif[0] > intervalo_trafo[1] and dif[0] <= intervalo_trafo[2]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Médio Risco')
                    elif dif[0] > intervalo_trafo[2] and dif[0] <= intervalo_trafo[3]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Alto Risco')
                    elif dif[0] > intervalo_trafo[3]:
                        texto = str(f'{name} | {dif[1]} fases detec | Class: Alto Risco Iminente')
                    caixa_text.insert(tk.INSERT, str(f'{texto}.\n'))
                    partes = texto.split(' | ')
                    equip_analis.append([name,dif[0],dif[1],partes[2].upper(),'trafo'])
            
    caixa_text.config(state=tk.DISABLED)
    return equip_analis

def calculo_medias(chaves):

    somas = {}
    ocorrencias = {}
    for elemento in chaves:
        chave = (elemento[0], elemento[1])
        valor = elemento[2]
        if chave in somas:
            somas[chave] += valor
            ocorrencias[chave] += 1
        else:
            somas[chave] = valor
            ocorrencias[chave] = 1
    medias = {}
    for chave, soma in somas.items():
        ocorrencia = ocorrencias[chave]
        media = soma / ocorrencia
        medias[chave] = media
    return medias

def trifasico(chaves):
    
    nomes = {}
    diferencas = {}
    for nomefase, tempmax in chaves.items():
        nome = nomefase[0]
        fase = nomefase[1]
        temp = tempmax.round(2)
        if nome in nomes:
            nomes[nome] += (fase + "-" + str(temp),)
        else:
            nomes[nome] = (fase + "-" + str(temp),)
        nomes[nome] = sorted(nomes[nome])
         
    for name, temp in nomes.items():
        qfase = len(temp)
        if qfase == 1:
            diferencas[name] = (0,1)
        elif qfase == 2:
            temp1 = float(temp[0].split('-')[1])
            temp2 = float(temp[1].split('-')[1])
            dif = round(abs(temp2-temp1),2)
            diferencas[name] = (dif,2)
        elif qfase == 3:
            temp1 = float(temp[0].split('-')[1])
            temp2 = float(temp[1].split('-')[1])
            temp3 = float(temp[2].split('-')[1])
            lista = [round(abs(temp2-temp1),2), round(abs(temp3-temp1),2), round(abs(temp3-temp2),2)]
            dif = max(lista, key=int)
            diferencas[name] = (dif,3)
    return diferencas      

#=====================================================SELECIONAR IMAGENS===========================================================#

def selec_imagens(lista1):
    
    chaves_mono = []
    trafos_mono = []
    
    for i in range(len(img_chaves)):
        for j in range(len(box_chaves[i])):
            nome, _ = equipamento(img_chaves[i], 0, j)
            tmax, _, _, _, _, _, _ = info_temp(img_chaves[i], 0, j)
            box_aux = [nome, img_chaves[i], tmax]
            chaves_mono.append(box_aux)
            
    for i in range(len(img_trafos)):
        for j in range(len(box_trafos[i])):
            nome, _ = equipamento(img_trafos[i], 1, j)
            tmax, _, _, _, _, _, _ = info_temp(img_trafos[i], 1, j)
            box_aux = [nome, img_trafos[i], tmax]
            trafos_mono.append(box_aux)
            
    lista2 = chaves_mono + trafos_mono
    
    lista_equip_imag = []
    lista_equip_tmax = []
    
    for i in range(len(lista1)):
        
        box_aux1 = []
        box_aux2 = []
        
        for j in range(len(lista2)):
            if lista1[i] == lista2[j][0]:
                box_aux1.append(lista2[j][1])
                box_aux2.append(lista2[j][2])
        
        box_aux3 = []
        for item in box_aux1:
            if item not in box_aux3:
                box_aux3.append(item)

        qtd = len(box_aux3)
        
        if qtd >= 5:
            img1 = box_aux3[2]
            img2 = box_aux3[qtd-2]
        elif qtd == 4 or qtd == 3:
            img1 = box_aux3[1]
            img2 = box_aux3[2]
        elif 3 > qtd:
            img1 = box_aux3[0]
            img2 = box_aux3[0]
        
        tmaximo = max(box_aux2)
        Tmaximo = float(round(tmaximo,2))
        
        lista_equip_imag.append([lista1[i],img1,img2])
        lista_equip_tmax.append([lista1[i],Tmaximo])
    
    return lista_equip_imag, lista_equip_tmax
        
#=======================================================EXIBIR IMAGENS=============================================================#

imagem_atual = ""

def exibir1(imagem):

    caminho_imagem = os.path.abspath(str(f'images/{imagem}'))
    image = cv2.imread(caminho_imagem)
    if combo.get() == 'Chave Seccionadora':
        item = img_chaves.index(imagem)
        box_aux = box_chaves[item]
        classe = 0
    elif combo.get() == 'Transformador':
        item = img_trafos.index(imagem)
        box_aux = box_trafos[item]
        classe = 1
    for i in range(len(box_aux)):
        x, y, w, h = box_aux[i]
        x, y, w, h = limite_coord(x,y,w,h)
        cv2.rectangle(image, (x,y), (x+w,y+h), (0,255,0), 1)
        nome, fase = equipamento(imagem, classe, i)
        fase = str(f'Fase {fase}')
        cv2.putText(image, nome, (x,y-7), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0,255,0), 1)
        cv2.putText(image, fase, (x,y+h+15), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0,255,0), 1)
        # TEMPERATURA
        tmax, tmin, _, ymax, xmax, ymin, xmin = info_temp(imagem, classe, i)
        verticesx = np.array([[xmax, ymax], [xmax-4, ymax+4], [xmax+4, ymax+4]], np.int32)
        verticesx = verticesx.reshape((-1, 1, 2))
        cv2.fillPoly(image, [verticesx], (100,255,0))
        #verticesn = np.array([[xmin, ymin], [xmin-4, ymin+4], [xmin+4, ymin+4]], np.int32)
        #verticesn = verticesn.reshape((-1, 1, 2))
        #cv2.fillPoly(image, [verticesn], (0,255,0))
        texto_max = str(f'{tmax:.2f}'); #texto_min = str(f'{tmin:.2f}')
        cv2.putText(image, texto_max, (xmax+7,ymax+4), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (100,255,0), 1)
        #cv2.putText(image, texto_min, (xmin+7,ymin+4), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0,255,0), 1)
    
    return image

def exibir(imagem):

    global imagem_atual
    imagem_atual = imagem
    image = exibir1(imagem)
    imagem_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    imagem_cv = ImageTk.PhotoImage(Image.fromarray(imagem_rgb))    
    rotulo.config(image=imagem_cv)
    rotulo.image = imagem_cv
    text_imagem["text"] = imagem_atual

def visualizar():

    if img_chaves == [] and img_trafos == []:
        caixa_text.config(state=tk.NORMAL)
        caixa_text.insert(tk.INSERT, 'Sem equipamento detectado.\n')
        caixa_text.config(state=tk.DISABLED)
    else:
        if combo.get() == 'Selecione':
            caixa_text.config(state=tk.NORMAL)
            caixa_text.insert(tk.INSERT, 'Selecione um equipamento para exibir.\n')
            caixa_text.config(state=tk.DISABLED)
        elif combo2.get() == 'Selecione':
            if combo.get() == 'Chave Seccionadora':
                exibir(img_chaves[0])
            elif combo.get() == 'Transformador':
                exibir(img_trafos[0])
        else:
            exibir(combo2.get())

def proxima_imagem():

    if imagem_atual != "":
        if combo.get() == 'Chave Seccionadora':
            indice = img_chaves.index(imagem_atual)
            if indice+1 != len(img_chaves):
                exibir(img_chaves[indice+1])
                
        elif combo.get() == 'Transformador':
            indice = img_trafos.index(imagem_atual)
            if indice+1 != len(img_trafos):
                exibir(img_trafos[indice+1])

def anterior_imagem():

    if imagem_atual != "":
        if combo.get() == 'Chave Seccionadora':
            indice = img_chaves.index(imagem_atual)
            if indice != 0:
                exibir(img_chaves[indice-1])      
        elif combo.get() == 'Transformador':
            indice = img_trafos.index(imagem_atual)
            if indice != 0:
                exibir(img_trafos[indice-1])


#=======================================================SALVAR IMAGENS=============================================================#

def salvar_imagem():

    if imagem_atual != "":
        janela_carregar = tk.Tk()
        janela_carregar.withdraw()
        tipos_arquivo = [("Imagens PNG", "*.png"), ("Imagens JPG", "*.jpg;*.jpeg"), ("Imagens BMP", "*.bmp")]
        caminho = filedialog.asksaveasfilename(defaultextension=".png", filetypes=tipos_arquivo)
        
        # Se o usuário cancelar, sai da função
        if not caminho:
            janela_carregar.destroy()
            return
        
        image = exibir1(imagem_atual)
        cv2.imwrite(caminho, image)
        janela_carregar.destroy()


#=====================================================INTERFACE DE USUÁRIO=========================================================#

aux_r = True; aux_c = True

janela = tk.Tk()
janela.title("UAV Thermal Scout")
janela.resizable(False, False)
janela.iconbitmap(os.path.abspath('files/icones/eletrobras.ico'))

subframe1 = tk.Frame(janela)
subframe1.grid(column=0,row=0)

# Texto de título
texto_exe = tk.Label(subframe1, text="Selecione um intervalo (Exemplo: 1-5):")
texto_exe.grid(column=0, row=0, pady=10)

texto_intervalo = tk.Entry(subframe1)
texto_intervalo.grid(column=2,row=0)

# Tela de log
caixa_text = scrolledtext.ScrolledText(janela, wrap=tk.WORD, width = 70, height = 31)
caixa_text.grid(column=0, row=1, padx=10, pady=0)

# Botão de iniciar processamento
botao_ini = tk.Button(janela, text="INICIAR DETECÇÃO DE EQUIPAMENTOS", command=process_imagens)
botao_ini.grid(column=1, row=0, pady=10)

# Exibição de imagem
imagem_branco = Image.new("RGB",(640,512),"white")
imagem_tk = ImageTk.PhotoImage(imagem_branco)
rotulo = tk.Label(janela, image=imagem_tk)
rotulo.grid(column=1, row=1, padx=10, pady=0)

#======================================================================#

subframe2 = tk.Frame(janela)
subframe2.grid(column=0, row=2, rowspan=2, columnspan=1)

# Selecionar tipo de equipamento
text_combo = tk.Label(subframe2, text="Selecione o equipamento:")
text_combo.grid(column=0, row=0)
itens = ['Chave Seccionadora','Transformador']
combo = ttk.Combobox(subframe2, values=itens, state="readonly")
combo.grid(column=0, row=1)
combo.set('Selecione')
combo.bind("<<ComboboxSelected>>", select_combo)

# Selecionar imagem
text_combo2 = tk.Label(subframe2, text="Selecione a imagem:")
text_combo2.grid(column=1, row=0)
combo2 = ttk.Combobox(subframe2, state="readonly")
combo2.grid(column=1, row=1, padx=20)
combo2.set('Selecione')

# Botão visualizar
botao_ver = tk.Button(subframe2, text="VISUALIZAR", command=visualizar)
botao_ver.grid(column=2,row=1, padx=(40,0), pady=(0,10))

#======================================================================#

# Sub janela
subframe3 = tk.Frame(janela)
subframe3.grid(column=1, row=2, rowspan=2, columnspan=1)

botao_anterior = tk.Button(subframe3, text="<", command=anterior_imagem )
botao_anterior.grid(column=0, row=0)
botao_proximo = tk.Button(subframe3, text=">", command=proxima_imagem)
botao_proximo.grid(column=2,row=0)

# Filtro de análise
text_imagem = tk.Label(subframe3, text="______________")
text_imagem.grid(column=1, row=0, padx=200, pady=(15,0))

#======================================================================#

subframe4 = tk.Frame(janela)
subframe4.grid(column=2,row=0,rowspan=5, padx=(0,10))

# Botão limpar
borracha = tk.PhotoImage(file=os.path.abspath('files/icones/borracha.png'))
botao_limpar = tk.Button(subframe4, image=borracha, command=limpar)
botao_limpar.grid(column=0, row=0, pady=40)

# Botão detectar
lupa = tk.PhotoImage(file=os.path.abspath('files/icones/lupa.png'))
botao_detector = tk.Button(subframe4, image=lupa, command=detector)
botao_detector.grid(column=0, row=1)

# Botão salvar imagem
salvar = tk.PhotoImage(file=os.path.abspath('files/icones/salvar.png'))
botao_salvar = tk.Button(subframe4, image=salvar, command=salvar_imagem)
botao_salvar.grid(column=0, row=2, pady=40)

#=================================PARÂMETROS DO SOFTWARE=====================================#

def configuracao():

    global aux_c

    if aux_c == True:

        aux_c = False

        janela2 = tk.Tk()
        janela2.title("Configurações")
        janela2.resizable(False, False) # Não deixa maximizar a janela
        janela2.iconbitmap(os.path.abspath('files/icones/eletrobras.ico'))

        subjanela = tk.Frame(janela2)
        subjanela.grid(column=0,row=0,columnspan=2,padx=90,pady=(30,0))

        texto_gsd = tk.Label(subjanela,text="GSD: ")
        texto_gsd.grid(column=0,row=0)
        entry_gsd = tk.Entry(subjanela); entry_gsd.insert(0,str(gsd_imagem))
        entry_gsd.grid(column=1,row=0)

        texto_int_chave = tk.Label(subjanela,text="IRC: ")
        texto_int_chave.grid(column=0,row=1)
        entry_int_chave = tk.Entry(subjanela); entry_int_chave.insert(0,str(intervalo_chave))
        entry_int_chave.grid(column=1,row=1)

        texto_int_trafo = tk.Label(subjanela,text="IRT: ")
        texto_int_trafo.grid(column=0,row=2)
        entry_int_trafo = tk.Entry(subjanela); entry_int_trafo.insert(0,str(intervalo_trafo))
        entry_int_trafo.grid(column=1,row=2)

        texto_filtro = tk.Label(subjanela,text="FCG: ")
        texto_filtro.grid(column=0,row=3)
        entry_filtro = tk.Entry(subjanela); entry_filtro.insert(0,str(filtro_deteccao))
        entry_filtro.grid(column=1,row=3)

        texto_confianca = tk.Label(subjanela,text="CONF: ")
        texto_confianca.grid(column=0,row=4)
        entry_confianca = tk.Entry(subjanela); entry_confianca.insert(0,str(confianca))
        entry_confianca.grid(column=1,row=4)

        def setparametros():
            global gsd_imagem, intervalo_chave, aux_c
            global intervalo_trafo, filtro_deteccao, confianca
            try:
                gsd_imagem = float(entry_gsd.get())
                intervalo_chave = ast.literal_eval(entry_int_chave.get())
                intervalo_trafo = ast.literal_eval(entry_int_trafo.get())
                filtro_deteccao = float(entry_filtro.get())
                confianca = float(entry_confianca.get())
                with open(caminho_txt, 'w') as arquivo:
                    arquivo.write(f'gsd_imagem={gsd_imagem}\n')
                    arquivo.write(f'intervalo_chave={intervalo_chave}\n')
                    arquivo.write(f'intervalo_trafo={intervalo_trafo}\n')
                    arquivo.write(f'filtro_deteccao={filtro_deteccao}\n')
                    arquivo.write(f'confianca={confianca}')
            except Exception as e:
                messagebox.showerror("Erro de Execução", str(e))
            aux_c = True
            janela2.destroy()

        subjanela2 = tk.Frame(janela2)
        subjanela2.grid(column=0,row=1,columnspan=2)
        botao_ok = tk.Button(subjanela2, text="          OK          ", command=setparametros)
        botao_ok.grid(column=0,row=1,pady=30)

        def aux_c_fechar():
            global aux_c
            aux_c = True
            janela2.destroy()

        janela2.protocol("WM_DELETE_WINDOW", aux_c_fechar)

# Botão Configurações
config = tk.PhotoImage(file=os.path.abspath('files/icones/engrenagem.png'))
botao_config = tk.Button(subframe4, image=config, command=configuracao)
botao_config.grid(column=0, row=4, pady=40)

#================================== GERAR RELATÓRIO ====================================#

def janela_relatorio():

    global aux_r

    if aux_r == True:

        aux_r = False
    
        relatorio_janela = tk.Tk()
        relatorio_janela.title("Relatório")
        relatorio_janela.resizable(False, False) # Não deixa maximizar a janela
        relatorio_janela.iconbitmap(os.path.abspath('files/icones/eletrobras.ico'))

        subjanela = tk.Frame(relatorio_janela)
        subjanela.grid(column=0,row=0,columnspan=2,padx=90,pady=(30,0))

        texto_titulo = tk.Label(subjanela,text="TÍTULO: ")
        texto_titulo.grid(column=0,row=0)
        entry_titulo = tk.Entry(subjanela, width=60)
        entry_titulo.grid(column=1,row=0)

        texto_autor = tk.Label(subjanela,text="AUTOR: ")
        texto_autor.grid(column=0,row=1)
        entry_autor = tk.Entry(subjanela, width=60)
        entry_autor.grid(column=1,row=1)
        
        texto_observ = tk.Label(subjanela,text="OBSERVAÇÕES: ")
        texto_observ.grid(column=0,row=2)
        entry_observ = tk.Entry(subjanela, width=60)
        entry_observ.grid(column=1,row=2)
        
        # ========================================== CONFECÇÃO DO RELATÓRIO ======================================= #
        
        def conf_relatorio(Equip_Analisados):
            
            # Aquisição de Informações para o Relatório
            subtitulo = entry_titulo.get()
            autor = entry_autor.get()
            observ = entry_observ.get()
            
            prim_imagem = Image.open(str(f'images/{nomes_imagens[0]}'))
            ultm_imagem = Image.open(str(f'images/{nomes_imagens[quant_imagens-1]}'))
        
            exif_dados = prim_imagem._getexif()
            for tag, valor in exif_dados.items():
                tag_nome = TAGS.get(tag, tag)
                if tag_nome == "DateTime":
                    datahora1 = valor
                    
            exif_dados = ultm_imagem._getexif()
            for tag, valor in exif_dados.items():
                tag_nome = TAGS.get(tag, tag)
                if tag_nome == "DateTime":
                    datahora2 = valor
        
            datahora1 = datetime.strptime(datahora1, '%Y:%m:%d %H:%M:%S')
            datahora2 = datetime.strptime(datahora2, '%Y:%m:%d %H:%M:%S')
        
            data_inspec = datahora1.strftime('%d/%m/%Y')
            hora_inicio = datahora1.strftime('%H:%M')
            hora_termin = datahora2.strftime('%H:%M')
            
            for tag, valor in exif_dados.items():
                tag_nome = TAGS.get(tag, tag)
                if tag_nome == "Model":
                    Modelo = valor
            Modelo = Modelo.replace('-',' ')
            
            Imagens = list(set(img_chaves + img_trafos))
            
            Equipamentos = []
            for i in range(len(Equip_Analisados)):
                Equipamentos.append(Equip_Analisados[i][0])
            
            # Informações para seção "Diagnósticos"
            # Lista de equipamentos:
            
            Imagens_Selecionadas, Temp_Maxima = selec_imagens(Equipamentos)
                
            # Gerar Documento
            doc = Document()
            
            doc.add_picture("files/icones/logo.png")
            imagem_paragrafo = doc.paragraphs[-1]
            imagem_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar título
            titulo_doc = doc.add_paragraph()
            titulo_run = titulo_doc.add_run("UAV THERMAL SCOUT")
            titulo_run.bold = True
            titulo_run.font.size = Pt(14)
            titulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar Subtítulo
            subtitulo_doc = doc.add_paragraph()
            subtitulo_run = subtitulo_doc.add_run(subtitulo)
            subtitulo_run.bold = True
            subtitulo_run.font.size = Pt(12)
            subtitulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar Autor
            autor_doc = doc.add_paragraph()
            autor_doc.add_run(str(f'POR: {autor}'))
            autor_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar Datas e Horas
            doc.add_paragraph(str(f'DATA DA INSPEÇÃO: {data_inspec}    |    INÍCIO: {hora_inicio}    |    TÉRMINO: {hora_termin}'))
            doc.add_paragraph(str(f'DRONE UTILIZADO: {Modelo}'))
            
            # Adicionar seção 'Imagens Analisadas'
            secao_img = doc.add_paragraph()
            secao_img_run = secao_img.add_run("| IMAGENS ANALISADAS:")
            secao_img_run.bold = True
            
            tabela = doc.add_table(rows=1, cols=5)
            row = tabela.rows[0].cells
            
            for i, img in enumerate(Imagens):
                row[i % 5].text = img
                if (i + 1) % 5 == 0:
                    row = tabela.add_row().cells
            
            # Adicionar seção 'Equipamentos Analisados'
            secao_equi = doc.add_paragraph()
            secao_equi_run = secao_equi.add_run("| EQUIPAMENTOS ANALISADOS:")
            secao_equi_run.bold = True
            
            tabela2 = doc.add_table(rows=1, cols=5)
            row2 = tabela2.rows[0].cells
            
            for i, equi in enumerate(Equipamentos):
                row2[i % 5].text = equi
                if (i + 1) % 5 == 0:
                    row2 = tabela2.add_row().cells
            
            # Adicionar seção 'Diagnósticos'
            secao_diag = doc.add_paragraph()
            secao_diag_run = secao_diag.add_run("| DIAGNÓSTICOS:")
            secao_diag_run.bold = True
            
            for i in range(len(Equipamentos)):
                nome_op = Equip_Analisados[i][0]
                dif_tem = Equip_Analisados[i][1]
                fases_d = Equip_Analisados[i][2]
                class_r = Equip_Analisados[i][3]
                temp_ma = Temp_Maxima[i][1]
                imag_t1 = Imagens_Selecionadas[i][1]
                imag_t2 = Imagens_Selecionadas[i][2]
                selecao = Equip_Analisados[i][4]

                texto_nome_op = doc.add_paragraph()
                texto_nome_run = texto_nome_op.add_run(str(f'# {nome_op}'))
                texto_nome_run.bold = True

                if selecao == 'chave':
                    combo.set('Chave Seccionadora')
                    tipo_ep = 'SECCIONADORA'
                elif selecao == 'trafo':
                    combo.set('Transformador')
                    tipo_ep = 'TRANSFORMADOR'
                
                texto_dia = str(f'| {tipo_ep} | FASES IDENT.: {fases_d} | DIF. TEMP.: {dif_tem}°C | MÁX. TEMP.: {temp_ma}°C |')           
                doc.add_paragraph(texto_dia)
                                
                # Adição das Imagens no documento
                
                if selecao == 'chave':
                    combo.set('Chave Seccionadora')
                elif selecao == 'trafo':
                    combo.set('Transformador')

                imagem1_np = exibir1(imag_t1)
                imagem1 = cv2.cvtColor(imagem1_np, cv2.COLOR_BGR2RGB)
                buffer1 = BytesIO()
                imagem1_pil = Image.fromarray(imagem1)
                imagem1_pil.save(buffer1, format="PNG")
                buffer1.seek(0)

                imagem2_np = exibir1(imag_t2)
                imagem2 = cv2.cvtColor(imagem2_np, cv2.COLOR_BGR2RGB)
                buffer2 = BytesIO()
                imagem2_pil = Image.fromarray(imagem2)
                imagem2_pil.save(buffer2, format="PNG")
                buffer2.seek(0)
                
                tabela_img = doc.add_table(rows=1, cols=2)
                
                cell_1 = tabela_img.cell(0, 0)
                run_01 = cell_1.paragraphs[0].add_run()
                run_01.add_picture(buffer1, width=Cm(7))
                
                cell_2 = tabela_img.cell(0, 1)
                run_02 = cell_2.paragraphs[0].add_run()
                run_02.add_picture(buffer2, width=Cm(7))

                buffer1.close()
                buffer2.close()

                texto_class = doc.add_paragraph()
                texto_class.add_run(str(f'{class_r}'))
                texto_class.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("---------------------------------------------------------------------------------------------------------------------")

            combo.set('Selecione')

            # Adicionar seção 'OBSERVAÇÕES'
            secao_equi = doc.add_paragraph()
            secao_equi_run = secao_equi.add_run("| OBSERVAÇÕES:")
            secao_equi_run.bold = True

            doc.add_paragraph(str(f'{observ}'))

            # ===== SALVAR RELATÓRIO =====#
            
            janela_salvar_relatorio = tk.Tk()
            janela_salvar_relatorio.withdraw()
            
            caminho_salvar = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
            
            if not caminho_salvar:
                janela_salvar_relatorio.destroy()
                return
            
            doc.save(caminho_salvar)
            
            janela_salvar_relatorio.destroy()
        
        # ======================================================================================================== #

        def setrelatorio():
            chaves_analis = []
            trafos_analis = []
            global aux_r
            try:
                if img_chaves == [] and img_trafos == []: 
                    process_imagens()
                    if img_chaves == [] and img_trafos == []:
                        messagebox.showerror("Erro de Execução", "Sem detecções!")
                        aux_r_fechar()
                    if img_chaves != []:
                        combo.set('Chave Seccionadora')
                        chaves_analis = detector()
                    if img_trafos != []:
                        combo.set('Transformador')
                        trafos_analis = detector()
                elif img_chaves != [] and img_trafos == []:
                    combo.set('Chave Seccionadora')
                    chaves_analis = detector()
                elif img_chaves == [] and img_trafos != []:
                    combo.set('Transformador')
                    trafos_analis = detector()
                else:
                    combo.set('Chave Seccionadora')
                    chaves_analis = detector()
                    combo.set('Transformador')
                    trafos_analis = detector()
                combo.set('Selecione')
                equip_analis = chaves_analis+trafos_analis
                conf_relatorio(equip_analis)
            except Exception as e:
                messagebox.showerror("Erro de Execução", str(e))
            aux_r_fechar()

        subjanela2 = tk.Frame(relatorio_janela)
        subjanela2.grid(column=0,row=1,columnspan=2)
        botao_ok = tk.Button(subjanela2, text="       GERAR RELATÓRIO       ", command=setrelatorio)
        botao_ok.grid(column=0,row=1,pady=30)

        def aux_r_fechar():
            global aux_r
            aux_r = True
            relatorio_janela.destroy()

        relatorio_janela.protocol("WM_DELETE_WINDOW", aux_r_fechar)

# Botão Relatório
relat = tk.PhotoImage(file=os.path.abspath('files/icones/relatorio.png'))
botao_relatorio = tk.Button(subframe4, image=relat, command=janela_relatorio)
botao_relatorio.grid(column=0, row=3)

janela.mainloop()

#print('Finalizado!')

